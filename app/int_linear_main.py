# -*- coding: utf-8 -*-
from tkinter import messagebox
from tkinter.filedialog import asksaveasfilename
from anytree import Node
from anytree.exporter import DotExporter
from docx import Document
from docx.shared import Cm
from openpyxl import load_workbook
from pulp import *
from math import ceil, floor
import string
from sympy import Symbol, integrate
import platform
import os
import matplotlib.pyplot as plt

'''глобальные переменные'''
# чтобы прописать graphviz в PATH
os.environ["PATH"] += os.pathsep + r'C:\Program Files (x86)\Graphviz2.38\bin'
# допустимые символы в массивах excel
expr_allowed_symbols = set(string.digits + '.' + '*' + '/' + '(' + ')' + 'x' + 'X')
symbols = set('.' + '*' + '/' + '(' + ')')


class Solved(object):

    def __init__(self, problem, number, func_value, vars_value, status,
                 cont_var=None, cont_var_value=None, parent_number=None, e=None):
        self.problem = problem
        self.status = status
        self.number = number
        self.func_value = func_value
        self.vars_value = vars_value
        self.cont_var = cont_var
        self.cont_var_value = cont_var_value
        self.parent_number = parent_number
        self.e = e
        self.make_node()

    def __repr__(self):
        return str(self.number)

    # чтобы решить проблему с pyinstaller
    if platform.system() == 'Windows':
        solver = COIN_CMD(path=os.path.join(os.getcwd(), 'solver\\win\\cbc.exe'))
    else:
        solver = COIN_CMD(path=os.path.join(os.getcwd(), 'solver/linux/cbc'))
    tree = []
    statuses = ['Не решено', 'Оптимально', 'Неопределенно', 'Не ограниченно', 'Нерешаемо']

    def nodenamefunc(node):
        node_params = [str(node.name), node.status, str(node.func_value), *node.xs]
        return "\n".join(node_params)

    def make_node(self):
        status = Solved.statuses[self.status]
        xs = [str(x[0]) + ' = ' + str(x[1]) for x in zip(self.problem.variables(), self.vars_value)]
        new_node = Node(name=self.number, status=status, xs=xs,
                        func_value=self.func_value, parent_name=self.parent_number)
        for node in self.tree:
            if node.name == new_node.parent_name:
                new_node.parent = node
        self.tree.append(new_node)


class Solution(object):
    def __init__(self, acc, solution=None, optimal_problems=[], auto_coeff_D=False):
        self.acc = acc
        self.coeff_D = auto_coeff_D
        self.has_sol = False
        if solution is not None:
            for x in solution.vars_value:
                if x != 0:
                    self.has_sol = True
                    self.optimal_problems = optimal_problems
                    self.variables = solution.problem.variables()
                    self.func_value = solution.func_value
                    self.vars_value = solution.vars_value
                    self.number = solution.number
                    break

    def __repr__(self):
        return str(self.number)


def create_Solved(problem, acc, parent_number=None):
    problem_copy = problem.deepcopy()
    problem_copy.solve(Solved.solver)
    acc += 1
    # создаем объект решенной задачи
    solved = Solved(problem=problem,
                    status=problem_copy.status,
                    number=acc,
                    func_value=value(problem_copy.objective),
                    vars_value=[var.varValue for var in problem_copy.variables()],
                    parent_number=parent_number)
    for v in problem_copy.variables():
        if v.varValue != int(v.varValue):
            solved.cont_var, solved.cont_var_value = v, v.varValue
            break
    return solved, acc


def get_values(worksheet, row, column, column_len=-1, is_expression=False):
    values = []
    allowed = expr_allowed_symbols
    while worksheet.cell(row, column).value is not None:
        #  если это выражение, все его символы разрешены и оно состоит не из одних лишь операторов
        cell_value = worksheet.cell(row, column).value
        if is_expression and set(str(cell_value)) <= allowed and not set(str(cell_value)) <= symbols:
            values.append(str(cell_value))
            row += 1
        elif (type(cell_value) is float) or (type(cell_value) is int):
            values.append(cell_value)
            row += 1
        else:
            raise Exception(
                messagebox.showinfo('Ошибка',
                                    'В листе Excel в ячейке (' + str(column) + ', ' + str(row) + ') '
                                    'значение написано неверно или использованы недопустимые символы. '
                                    '\nДопустимые: ' + ' '.join(sorted(allowed)))
            )
    if column_len != -1 and len(values) != column_len:
        raise Exception(
            messagebox.showinfo('Ошибка', 'В листе Excel кол-во элементов в колонке '
                                + str(column) + ' не равно кол-ву элементов в предшествующей'))
    return values  # не прерывается


def get_data_excel(filepath, T):
    #  открываем нужный лист в выбранной в форме книге excel
    workbook = load_workbook(filename=str(filepath))
    worksheet = workbook.worksheets[0]
    #  получаем массивы данных с листа по столбцам
    alpha = get_values(worksheet, 2, 2)
    beta = get_values(worksheet, 2, 3, len(alpha))
    v = get_values(worksheet, 2, 4, len(alpha))
    V = get_values(worksheet, 2, 5, len(alpha))
    teta = get_values(worksheet, 2, 6, len(alpha), is_expression=True)
    # интегрируем тета
    x = Symbol('x')
    try:
        teta_integrated = [float(integrate(eval(teta[i]), (x, 0, T))) for i in range(0, len(teta))]
    except Exception:
        messagebox.showinfo('Ошибка', 'Не все функции по тета интегрируются')
        raise
    data = {'alpha': alpha,
            'beta': beta,
            'v': v,
            'teta_integrated': teta_integrated,
            'V': V}
    return data, workbook, worksheet


def sort_data(sort_type, data):
    if sort_type == 'b/a':
        ba = [b / a for a, b in zip(data['alpha'], data['beta'])]
        zipped = list(zip(ba, data['teta_integrated'], data['alpha'],
                          data['beta'], data['V'], data['v']))
        zipped.sort(reverse=True)
        ba, data['teta_integrated'], data['alpha'], data['beta'], data['V'], data['v'] = zip(*zipped)
    elif sort_type == 'teta':
        zipped = list(zip(data['teta_integrated'], data['alpha'],
                          data['beta'], data['V'], data['v']))
        zipped.sort(reverse=True)
        data['teta_integrated'], data['alpha'], data['beta'], data['V'], data['v'] = zip(*zipped)
    return data


def form_problem(data, **coeffs):

    def form_problem_1(alpha, beta, v, teta_integrated, V, F):
        k = [a / b for a, b in zip(V, v)]
        n = len(alpha)
        problem = LpProblem('Zadachka', LpMaximize)
        x = LpVariable.dicts('x', range(n), lowBound=0, cat=LpContinuous)
        sum_var1 = lpSum([x[i] * v[i] * beta[i] for i in range(0, n)])
        sum_var2 = lpSum([x[i] * v[i] * alpha[i] for i in range(0, n)])
        problem += sum_var1 - sum_var2  # 'Функция цели "11.1"'
        problem += sum_var2 <= F  # "11.2"
        constraint1 = [x[i] <= k[i] for i in range(0, n)]
        for cnstr in constraint1:
            problem += cnstr
        constraint2 = [x[i] * v[i] <= min(teta_integrated[i], V[i]) for i in range(0, n)]
        for cnstr in constraint2:
            problem += cnstr
        constraint3 = [min(teta_integrated[i], V[i]) <= v[i] * (x[i] + 1) for i in range(0, n)]
        for cnstr in constraint3:
            problem += cnstr
        return problem

    def form_problem_2(alpha, beta, v, teta_integrated, V, F, D, y):
        k = [a / b for a, b in zip(V, v)]
        n = len(alpha)
        V = [v[i] * k[i] for i in range(0, n)]
        problem = LpProblem('Zadachka', LpMaximize)
        x = LpVariable.dicts('x', range(n), lowBound=0, cat=LpContinuous)
        sum_xvb = lpSum([x[i] * v[i] * beta[i] for i in range(0, n)])
        sum_xva = lpSum([x[i] * v[i] * alpha[i] for i in range(0, n)])
        sum_1yax = lpSum([(1 + y) * alpha[i] * x[i] for i in range(0, n)])
        sum_bx = lpSum([beta[i] * x[i] for i in range(0, n)])
        problem += sum_xvb + ((1 + y) * (F - sum_xva))  # цель
        problem += sum_xva <= F + D
        constraint1 = [x[i] <= k[i] for i in range(0, n)]
        constraint2 = [x[i] * v[i] <= min(teta_integrated[i], V[i]) for i in range(0, n)]
        constraint3 = [min(teta_integrated[i], V[i]) <= v[i] * (x[i] + 1) for i in range(0, n)]
        for cnstr in [*constraint1, *constraint2, *constraint3]:
            problem += cnstr
        problem += sum_1yax <= sum_bx
        return problem

    alpha = data['alpha']
    beta = data['beta']
    v = data['v']
    V = data['V']
    teta_integrated = data['teta_integrated']
    if coeffs['zadacha'] == 1:
        return form_problem_1(alpha, beta, v, teta_integrated, V, coeffs['F'])
    elif coeffs['zadacha'] == 2:
        return form_problem_2(alpha, beta, v, teta_integrated, V, coeffs['F'], coeffs['D'], coeffs['y'])


def solve_problem(problem):
    # обнуляем дерево, потому что в одном прогоне программы может решаться несколько задач
    Solved.tree = []
    optimal = []  # лист оптимальных целочисленных решений
    queue = []  # очередь на ветвление
    acc = 0  # счетчик задач
    max_z = 0  # максимальное значение целевой функции

    first, acc = create_Solved(problem, acc, parent_number=None)
    # возвращаем задачу, если она не оптимальна
    if first.status != 1:
        return Solution(acc=acc, solution=None)
    else:
        # проверяем проблему на целочисленность
        if first.cont_var is not None:
            # добавляем первую проблему в очередь на ветвление
            queue.append(first)
            return branch_and_bound(queue, max_z, acc, optimal)
        # если проблема целочислена на первом шаге
        else:
            return Solution(acc=acc, solution=first, optimal_problems=[first])


def branch_and_bound(queue, max_z, acc, optimal):  # передаем сюда задачу на ветвление, в том числе нецелую переменную
    if queue:
        max_prob = queue[0]
        for prob in queue[1:]:
            if prob.value > max_prob.value:
                max_prob = prob
        queue.remove(max_prob)
        i = 1
        queue, max_z, acc, optimal = make_branch(max_prob, acc, queue, max_z, optimal, i)
        return branch_and_bound(queue, max_z, acc, optimal)
    else:
        if optimal:
            solution = optimal[0]
            for prob in optimal[1:]:
                if prob.func_value > solution.func_value:
                    solution = prob
            return Solution(acc=acc, solution=solution, optimal_problems=optimal)
        else:
            return Solution(acc=acc, solution=None)


def make_branch(parent_problem, acc, queue, max_z, optimal, i):
    child_problem = parent_problem.problem.deepcopy()
    if i == 1:
        child_problem += parent_problem.cont_var <= floor(parent_problem.cont_var_value)  # левая ветвь
    elif i == 2:
        child_problem += parent_problem.cont_var >= ceil(parent_problem.cont_var_value)  # правая ветвь
    child_solved, acc = create_Solved(child_problem, acc, parent_number=parent_problem.number)
    if child_solved.status == 1:  # формируем подпроблему в очередь
        if child_solved.cont_var is None and max_z == 0:  # шаг 5
            max_z = child_solved.func_value
            optimal.append(child_solved)
            for prob in queue:  # проверяем, будем ли ветвить эту задачу
                if prob.value > max_z:
                    queue.append(child_solved)
                    break

        elif child_solved.cont_var is None and child_solved.func_value >= max_z:
            optimal.append(child_solved)

        elif child_solved.cont_var is not None and max_z == 0:
            queue.append(child_solved)

    if i == 1:
        return make_branch(parent_problem, acc, queue, max_z, optimal, i + 1)  # ветвим второй раз
    elif i == 2:
        return queue, max_z, acc, optimal


def solution_stability(solution, data):

    def minimum_e(optimal_list, e_min_problem, e_min_prob_index, es):
        """
        Эта функция находит интервалы значений инфляции (е), в рамках которых
        оптимальный портфель закупок сохраняется. Первый шаг расчета начинается с
        оптимальной проблемы, у которой е = 0.
        """
        e_min = float('inf')
        x_l = e_min_problem.vars_value
        l = e_min_prob_index
        n = len(optimal_list)
        for i in range(l + 1, n):
            x = optimal_list[i].vars_value
            e = (
                    (
                            sum(x_l[i] * (data['beta'][i] - data['alpha'][i]) for i in range(0, len(x))) -
                            sum(x[i] * (data['beta'][i] - data['alpha'][i]) for i in range(0, len(x)))
                    ) /
                    (
                            sum(x[i] * data['beta'][i] for i in range(0, len(x))) -
                            sum(x_l[i] * data['beta'][i] for i in range(0, len(x)))
                    )
            )
            if 0 < e < e_min:
                e_min = e
                e_min_problem = optimal_list[i]
                e_min_prob_index = i
        if l != n - 1 and e_min != float('inf'):
            es[e_min] = e_min_problem
            return minimum_e(optimal_list, e_min_problem, e_min_prob_index, es)
        else:
            return es

    def make_stability_plot(es, data):
        # нужно построить функцию от неизвестного e, имея две точки: значение функции при e = 0
        # и при e, пересекающем оптимальную функцию. По сути, нужно только посчитать
        # значение целевой функции в этой точке и построить графики. Нужно использовать
        # проблемы l+1
        alpha = data['alpha']
        v = data['v']
        n = len(alpha)

        x_opt = es[0].vars_value
        opt_number = es[0].number
        beta = [b * 2 for b in data['beta']]
        sum_var1 = sum([x_opt[i] * v[i] * beta[i] for i in range(0, n)])
        sum_var2 = sum([x_opt[i] * v[i] * alpha[i] for i in range(0, n)])
        opt_zero_f_val = es[0].func_value
        opt_one_f_val = sum_var1 - sum_var2
        fig, ax = plt.subplots(nrows=1, ncols=1)
        ax.plot([0, 1], [opt_zero_f_val, opt_one_f_val], label='Оптимальная задача: ' + str(opt_number))
        del es[0]
        if es:
            intersection_f_list = []
            zero_f_list = [problem.func_value for problem in es.values()]
            for e, problem in es.items():
                x = problem.vars_value
                beta = [b * (1 + e) for b in data['beta']]
                sum_var1 = sum([x[i] * v[i] * beta[i] for i in range(0, n)])
                sum_var2 = sum([x[i] * v[i] * alpha[i] for i in range(0, n)])
                intersection_f_list.append(sum_var1 - sum_var2)
            for zero_f, cross_f, e, problem in zip(zero_f_list, intersection_f_list, es, es.values()):
                ax.plot([0, e, 2*e], [zero_f, cross_f, cross_f*2],
                        label='Целочисленная задача: ' + str(problem.number))
                ax.plot([e, e], [0, cross_f], ':')

        ax.set_ylabel('Значение целевой функции от инфляции, F(e)')
        ax.set_xlabel('Инфляция в долях, e')
        ax.set_title('Устойчивость решения')
        ax.legend()
        ax.grid()
        plot_path = 'results/temp_plot.png'
        fig.savefig(plot_path, bbox_inches='tight')
        return plot_path, es

    optimal_list = solution.optimal_problems
    optimal_list.append(Solved(problem=solution.optimal_problems[0].problem,
                               number='a', func_value=100, vars_value=[10, 9, 13],
                               status=1))
    optimal_list.append(Solved(problem=solution.optimal_problems[0].problem,
                               number='b', func_value=100, vars_value=[7, 10, 10],
                               status=1))
    optimal_list.append(Solved(problem=solution.optimal_problems[0].problem,
                               number='c', func_value=200, vars_value=[10, 10, 7],
                               status=1))
    optimal_list.append(Solved(problem=solution.optimal_problems[0].problem,
                               number='d', func_value=300, vars_value=[9, 7, 11],
                               status=1))
    optimal_list.sort(key=lambda problem: sum(problem.vars_value[i] * data['beta'][i]
                                              for i in range(0, len(data['beta']))))

    sol_num = solution.number
    index_list = [x.number for x in optimal_list]
    sol_index = index_list.index(sol_num) if sol_num in index_list else None
    es = {0: solution}
    es = minimum_e(optimal_list, solution, sol_index, es)
    return make_stability_plot(es, data)


def show_results(sort_type, solved):
    if not solved.has_sol:
        status = 'Статус: Нерешаемо'
        acc = 'Кол-во решенных ЗЛП: ' + str(solved.acc)
        results = [status, acc]
        if solved.coeff_D:
            results.append('Подобранный D: ' + str(solved.coeff_D))
    else:
        status = 'Статус: Оптимально'
        xs = [str(x[0]) + ' = ' + str(x[1]) for x in zip(solved.variables, solved.vars_value)]
        number_of_optimal = 'Номер оптимальной задачи: ' + str(solved.number)
        func_value = 'Значение целевой функции: ' + str(solved.func_value)
        sort_type = 'Сортировка: ' + sort_type
        acc = 'Кол-во решенных ЗЛП: ' + str(solved.acc)
        results = [status, func_value, *xs, sort_type, number_of_optimal, acc]
        if solved.coeff_D:
            results.append('Подобранный D: ' + str(solved.coeff_D))
    tree_img_path = 'results/temp_tree.png'
    DotExporter(Solved.tree[0], nodenamefunc=Solved.nodenamefunc).to_picture(tree_img_path)
    return results, tree_img_path


def write_to_excel(workbook, worksheet, filepath, results):
    for row in worksheet['j2:j10']:
        for cell in row:
            cell.value = None
    for i in range(2, len(results)):
        worksheet.cell(i, 10).value = results[i-2]
    try:
        workbook.save(filepath)
    except PermissionError:
        messagebox.showinfo('Ошибка', 'Нет доступа к указанному файлу Excel.\n'
                                      'Возможно, Вы не закрыли этот файл.\n'
                                      'Результат записан не будет.')


def write_to_docx(problem, results, plot_img_path, tree_img_path, is_stable, es):
    file_name = asksaveasfilename(defaultextension='.docx',
                                  initialdir='results',
                                  filetypes=[('Word Document (.docx)', '.docx')])
    if file_name:
        document = Document()
        # создаем кортеж values of OrderedDict, преобразуем в лист и индексируем
        obj = str(list(problem.constraints.values())[0])
        constrs = map(str, list(problem.constraints.values())[1:])
        document.add_heading('Максимизируем целевую функцию:', 1)
        document.add_paragraph(str(obj) + '\n\nС ограничениями:\n' +
                               '\n'.join(constrs))
        document.add_paragraph("\n".join(results))
        document.add_heading('Дерево решений задачи:', 1)
        document.add_picture(tree_img_path, width=Cm(10))
        if is_stable:
            document.add_heading('Устойчивость решения:', 1)
            if plot_img_path is not None:
                document.add_picture(plot_img_path, width=Cm(10))
                document.add_paragraph('Интервалы сохранения значения оптимального портфеля закупок:\n' +
                                       '[0, ' + ''.join(str(e) + ', ' for e in es) + '∞)')
            else:
                document.add_paragraph('Задача не решаема')
        document.save(file_name)
        messagebox.showinfo('Файл создан', 'Файл с результатом был успешно сохранен!')
        os.startfile(file_name)


def integer_lp(filepath, **coeffs):
    plot_img_path = None
    is_stable = False
    data, workbook, worksheet = get_data_excel(filepath, coeffs['T'])
    sorted_data = sort_data(coeffs['sort'], data)
    # если автоподбор параметра D
    if 'auto_D' in coeffs:
        # решаем проблему с первым значением D и убираем его из списка
        coeffs['D'] = coeffs['auto_D'].pop(0)
        problem = form_problem(sorted_data, **coeffs)
        solution_problem = solve_problem(problem)
        solution_problem.coeff_D = coeffs['D']
        if not solution_problem.has_sol:
            solution_problem.func_value = 0
        for d in coeffs['auto_D']:
            coeffs['D'] = d
            problem = form_problem(sorted_data, **coeffs)
            solved_problem = solve_problem(problem)
            if solved_problem.has_sol and solved_problem.func_value > solution_problem.func_value:
                solution_problem = solved_problem
                solution_problem.coeff_D = d
    else:
        problem = form_problem(sorted_data, **coeffs)
        solution_problem = solve_problem(problem)
    # если стабильность решения и True
    if 'stable' in coeffs and coeffs['stable']:
        is_stable = True
        if solution_problem.has_sol:
            plot_img_path, es = solution_stability(solution_problem, sorted_data)
    results, tree_img_path = show_results(coeffs['sort'], solution_problem)
    write_to_excel(workbook, worksheet, filepath, results)
    try:
        answer = messagebox.askyesno('Решение', "\n".join(results) +
                                     '\n\nХотите сохранить подробный результат\nв DOCX файл?')
    except PermissionError:
        messagebox.showinfo('Ошибка', 'Доступ к выбранному файлу невозможен. Может быть,'
                                      'Вы пытаетесь переписать открытый файл')
    if answer:
        write_to_docx(problem, results, plot_img_path, tree_img_path, is_stable, es)
    # удаляю временные файлы, которые закидываются в doc с результатами, графика может не быть
    os.remove(tree_img_path)
    try:
        os.remove(plot_img_path)
    except TypeError:
        pass


def main():
    print(integer_lp('excel/Zadachka2.xlsx', T=1, F=30000, zadacha=1, sort='b/a', stable=1))


if __name__ == '__main__':
    main()
